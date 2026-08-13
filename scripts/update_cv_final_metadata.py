#!/usr/bin/env python3
"""Apply the approved privacy and publication-metadata updates to the CV."""

from copy import deepcopy
from pathlib import Path
import sys
import re

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
CV = ROOT / "assets" / "Caterina-Mauri-CV.docx"


def paragraph_with(fragment, document):
    matches = [
        Paragraph(element, document)
        for element in document.element.body.iter(qn("w:p"))
        if fragment in Paragraph(element, document).text
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {fragment!r}, found {len(matches)}")
    return matches[0]


def replace_text(paragraph, text):
    """Replace text while retaining the formatting of the first text run."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def insert_after(paragraph, text):
    clone = deepcopy(paragraph._element)
    for child in list(clone):
        if child.tag != qn("w:pPr"):
            clone.remove(child)
    paragraph._element.addnext(clone)
    inserted = paragraph._parent.add_paragraph()
    inserted._element.getparent().remove(inserted._element)
    inserted._element = clone
    inserted._p = clone
    replace_text(inserted, text)
    return inserted


def main():
    document = Document(CV)

    # Keep professional contact details, remove unnecessary private information.
    replace_text(paragraph_with("Personal details:", document), "Professional details:")
    for fragment in (
        "Born in Milan",
        "Italian citizenship",
        "Married, three children",
    ):
        delete_paragraph(paragraph_with(fragment, document))

    replacements = {
        "with Nicola Grandi, La tipologia linguistica.":
            "(2022) with Nicola Grandi, La tipologia linguistica. Unità e diversità nelle lingue del mondo. Roma: Carocci.",
        "Il corpus KIParla. Una risorsa per la didattica dell’oralità":
            "(2025), with Silvia Ballarè, Claudia Borghetti, Paolo Della Putta and Eleonora Zucchini. “Il corpus KIParla. Una risorsa per la didattica dell’oralità”. Educazione Linguistica. Language Education 14: 213–230. https://doi.org/10.30687/elle/2280-6792/2025/02/004",
        "Replicability all the way up":
            "(2025), with Andrea Sansò. “Replicability all the way up: commentary on ‘Replication and methodological robustness in quantitative typology’ by Becker and Guzmán Naranjo”. Linguistic Typology 29(3): 513–518. https://doi.org/10.1515/lingty-2025-0033",
        "Parlare italiano in contesto migratorio":
            "(accepted, in press), with Eleonora Zucchini and Silvia Ballarè. “Parlare italiano in contesto migratorio: il corpus Stra-ParlaBO”. Études romanes de Brno.",
        "At the crossroads of typology and language(s) in use":
            "(2025), with Silvia Ballarè and Simone Mattiola. “At the crossroads of typology and language(s) in use”. Linguistic Typology at the Crossroads 5(2): I–VII. https://doi.org/10.60923/issn.2785-0943/24199",
        "Object encoding in spoken language data and antipassives":
            "(2025), with Silvia Ballarè and Andrea Sansò. “Object encoding in spoken language data and antipassives”. Linguistic Typology at the Crossroads 5(2): 289–322. https://doi.org/10.60923/issn.2785-0943/21768",
        "Shaping reference through temporality":
            "(2026a), with Ludovica Pannitto. “Shaping reference through temporality: The case of Italian nominals”. Lingue e Linguaggio 25(1): 31–61. https://doi.org/10.1418/121389",
        "Temporal, aspectual, and modal evaluation":
            "(2026b), with Ludovica Pannitto. “Temporal, aspectual, and modal evaluation in nominal reference: The case of Italian”. Folia Linguistica, ahead of print: 1–31. https://doi.org/10.1515/flin-2025-0142",
        "Reuse by design":
            "(2026), with Ludovica Pannitto. “Reuse by design: a pivot-based architecture for the KIParla corpus of spoken Italian”. Journal of Open Humanities Data 12, article 81: 1–17. https://doi.org/10.5334/johd.527",
        "Subjunctive/indicative alternation with verba putandi":
            "(2026), with Silvia Ballarè. “Subjunctive/indicative alternation with verba putandi in spoken Italian”. Folia Linguistica 61: 1–33. https://doi.org/10.1515/flin-2024-0053",
        "KIPasti: a corpus of kitchen-table conversations":
            "(accepted, in press), with Silvia Ballarè and Eleonora Zucchini. “KIPasti: a corpus of kitchen-table conversations in Italian”. Cuadernos de Filología Italiana.",
        "Semi-structured interviews as spontaneous interaction":
            "(accepted, in press), with Silvia Ballarè and Eleonora Zucchini. “Semi-structured interviews as spontaneous interaction: insights from the ParlaBO module of KIParla”. E-languages.",
        "with Davide Garassino":
            "(submitted), with Davide Garassino. “Wh-in-situ and clefted wh-interrogatives in spoken Italian: A corpus-based exploration”. Submitted to Linguistics.",
        "with Eleonora Zucchini exemplification":
            "(accepted), with Eleonora Zucchini. “Esemplificare in interazione: costruzione di categorie tra parlanti L1 e L2”. Forthcoming in L’interazione comunicativa: relazioni, pratiche, prospettive, Studi AItLA.",
        "Introducing KIParla Forest: seeds":
            "(2025), with Ludovica Pannitto, Eleonora Zucchini, Silvia Ballarè, Cristina Bosco and Manuela Sanguinetti. “Introducing KIParla Forest: seeds for a UD annotation of interactional syntax”. Proceedings of the Eighth International Conference on Dependency Linguistics (Depling, SyntaxFest 2025), 54–73. https://cris.unibo.it/bitstream/11585/1022332/1/2025.depling-1.5.pdf",
        "Is Semi-Automatic Transcription Useful":
            "(2026), with Martina Simonotti, Ludovica Pannitto, Eleonora Zucchini and Silvia Ballarè. “Is Semi-Automatic Transcription Useful in Corpus Creation? Preliminary Considerations on the KIParla Corpus”. Proceedings of the Fifteenth Language Resources and Evaluation Conference (LREC 2026), 6547–6560. https://doi.org/10.63317/3gw6i7cbs5rc",
        "Coconstructions in spoken data":
            "(2026), with Ludovica Pannitto, Sylvain Kahane, Kaja Dobrovoljc, Elena Battaglia, Bruno Guillaume and Eleonora Zucchini. “Coconstructions in Spoken Data: UD Annotation Guidelines and First Results”. Proceedings of the Ninth Workshop on Universal Dependencies (UDW 2026), 60–75. https://cris.unibo.it/bitstream/11585/1069389/1/2026.udw-1.0.pdf",
        "Say Again? The Limits of Whisper":
            "(2026), with Martina Simonotti, Ludovica Pannitto, Adriano Ferraresi and Gabriele Carioli. “Say again? The limits of Whisper with conversation: A case study on the KIParla corpus”. Proceedings of SPEAKABLE @ LREC 2026, 16–30. https://cris.unibo.it/bitstream/11585/1069387/1/2026.speakable-1.0.pdf",
        "The pragmatic values of connective negation":
            "(2026, in press), with Chiara Gianollo. “The pragmatic values of connective negation with né in contemporary Italian”. In Mena Lafkioui & Johan van der Auwera (eds.), Connective Negation. Berlin: De Gruyter Brill.",
        "AiTLA, exemplification":
            "(accepted), with Eleonora Zucchini. “Esemplificare in interazione: costruzione di categorie tra parlanti L1 e L2”. Forthcoming in L’interazione comunicativa: relazioni, pratiche, prospettive, Studi AItLA.",
        "Stra-ParlaBO corpus":
            "Stra-ParlaBO corpus (with Silvia Ballarè and Eleonora Zucchini, KIParla module): https://doi.org/10.60760/unibo/stra-parlabo",
        "Lecture on Grammatica emergente":
            "Lecture on Grammatica emergente tra tipologia e discorso. Questioni teoriche e metodologiche. PhD program in ‘Studi letterari, linguistici e storici’, University of Salerno. Duration: 3 hours.",
        "TBA. Plenary talk at the International Conference “Mitigation":
            "Mirative responses: rejection and confirmation in common ground management. Plenary talk at the International Conference “Mitigation and Mirativity in CG Management”, University of Tübingen, Germany.",
        "Bologna, 30/3/2026.": "Bologna, 13/8/2026.",
    }

    for fragment, new_text in replacements.items():
        replace_text(paragraph_with(fragment, document), new_text)

    # Publications present in the synchronized Unibo archive but absent from the CV.
    lists_anchor = paragraph_with("List constructions: towards a unified account", document)
    insert_after(
        lists_anchor,
        "(2018), with Francesca Masini and Paola Pietrandrea. “Lists: description, delimitation, definition. A foreword”. Italian Journal of Linguistics 30(1): 41–48. https://doi.org/10.26346/1120-2726-115",
    )

    intro_anchor = paragraph_with("Le parti del discorso", document)
    intro = insert_after(
        intro_anchor,
        "(2022), with Nicola Grandi. “Introduzione”. In Nicola Grandi & Caterina Mauri (eds.), La tipologia linguistica. Unità e diversità nelle lingue del mondo, 11–15. Roma: Carocci.",
    )
    insert_after(
        intro,
        "(2025), with Silvia Ballarè. “Presentazione”. In Silvia Ballarè & Caterina Mauri (eds.), CLUB Working Papers in Linguistics, vol. 9, 6–7. Bologna: CLUB. https://doi.org/10.6092/unibo/amsacta/8647",
    )

    document.save(CV)
    print(f"Updated {CV}")


def supplement_recent_metadata():
    """Add remaining verified 2025–2026 metadata after the main update."""
    document = Document(CV)
    replacements = {
        "The mirative values of Italian altro che":
            "(2025), with Antonia Russo. “The mirative values of Italian altro che”. In Susana Rodríguez Rosique (ed.), Expressing Surprise at the Crossroads: Mirativity, Exclamativity and Cooptation in Romance Languages, 143–178. Berlin: Mouton De Gruyter. https://doi.org/10.1515/9783111386683-007",
        "with Mira Ariel, “Or cycles”":
            "(2025), with Mira Ariel. “‘Or’ cycles”. In Maj-Britt Mosegaard Hansen and Richard Waltereit (eds.), Cyclic Change in Grammar and Discourse, 220–246. Oxford: Oxford University Press. https://doi.org/10.1093/9780198940661.003.0010",
        "Understand in interaction: the rise":
            "(2025), with Silvia Ballarè. “To understand in interaction: the rise of epistemic and evidential constructions based on capire in spoken Italian”. In Karolina Grzech and Henrik Bergqvist (eds.), Expanding the Boundaries of Epistemicity, 31–68. Berlin/New York: De Gruyter Mouton. https://doi.org/10.1515/9783111516233-002",
        "Investigating ‘other’ in cross-linguistic perspective":
            "(2026), with Chiara Gianollo. “Investigating ‘Other’ in Cross-Linguistic Perspective: Theoretical and Methodological Challenges”. In Patrícia Amaral et al. (eds.), Other: Ambiguity, Constraints, and Change, 229–266. Leiden: Brill. https://doi.org/10.1163/9789004744196_009",
        "A diachronic view on non-verbal predication":
            "(2026), with Andrea Sansò. “A diachronic view on non-verbal predication”. In Pier Marco Bertinetto, Luca Ciucci & Denis Creissels (eds.), Non-verbal Predication in the World’s Languages: A Typological Survey, vol. 1, 57–86. Berlin/New York: De Gruyter Mouton. https://doi.org/10.1515/9783110730982-002",
    }
    for fragment, new_text in replacements.items():
        replace_text(paragraph_with(fragment, document), new_text)
    document.save(CV)
    print(f"Supplemented {CV}")


def cleanup_updated_entries():
    """Remove legacy hyperlink residue and the duplicated AItLA placeholder."""
    document = Document(CV)
    for element in document.element.body.iter(qn("w:p")):
        paragraph = Paragraph(element, document)
        text = paragraph.text
        cleaned = re.sub(r"(https?://\S+?)\1", r"\1", text)
        if cleaned != text:
            paragraph.text = cleaned

    aitla = [
        Paragraph(element, document)
        for element in document.element.body.iter(qn("w:p"))
        if "Esemplificare in interazione: costruzione di categorie tra parlanti L1 e L2" in Paragraph(element, document).text
    ]
    if len(aitla) == 2:
        delete_paragraph(aitla[0])
    elif len(aitla) != 1:
        raise RuntimeError(f"Expected one or two AItLA entries, found {len(aitla)}")

    placeholders = [
        Paragraph(element, document)
        for element in document.element.body.iter(qn("w:p"))
        if Paragraph(element, document).text.strip().lower() == "(submitted) nature kiparla"
    ]
    for paragraph in placeholders:
        delete_paragraph(paragraph)

    small_fixes = {
        "Miratività di secondo livello: quando la sorpresa Miratività di secondo livello":
            "Lecture on Miratività di secondo livello: quando la sorpresa riguarda le aspettative. PhD Program in ‘Studi letterari, linguistici e storici’, University of Salerno.",
        "Lingue in cambiamentoattraverso il tempo":
            "Lecture on Lingue in cambiamento attraverso il tempo e le società - dati e teorie in convergenza. PhD Program in ‘Studi filologici e linguistici sul patrimonio scritto e orale’, University of Bergamo.",
        "1 started in September 2026":
            "Dept. of Modern languages, literatures and cultures, Univ. of Bologna. Supervision of 7 PostDocs (4 are now tenured, 2 started in January 2026, 1 will start in September 2026).",
    }
    for fragment, new_text in small_fixes.items():
        matches = [
            Paragraph(element, document)
            for element in document.element.body.iter(qn("w:p"))
            if fragment in Paragraph(element, document).text
        ]
        if matches:
            replace_text(matches[0], new_text)

    # Fixed row heights caused long entries and dates to overlap in the PDF.
    for row in document.element.body.iter(qn("w:tr")):
        properties = row.find(qn("w:trPr"))
        if properties is not None:
            for height in list(properties.findall(qn("w:trHeight"))):
                properties.remove(height)

    document.save(CV)
    print(f"Cleaned {CV}")


if __name__ == "__main__":
    if "--supplement" in sys.argv:
        supplement_recent_metadata()
    elif "--clean" in sys.argv:
        cleanup_updated_entries()
    else:
        main()
