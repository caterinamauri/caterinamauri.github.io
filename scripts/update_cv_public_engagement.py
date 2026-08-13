from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path("/Users/caterinamauri/Library/CloudStorage/OneDrive-AlmaMaterStudiorumUniversitàdiBologna/A_mywork/CV/Mauri_CV.docx")
OUTPUT = Path("assets/Caterina-Mauri-CV.docx")


ENTRIES = [
    ("2024-now", "TransformAzioni: Officina delle lingue, letterature e culture in transito. Contributor to the LILEC public-engagement project and to the participatory workshop ‘AttraVerso la mia città: lingue, parole e voci a Bologna’, working with local communities on language, belonging, migration and urban space.", "https://site.unibo.it/transformazioni-lingue-letterature-culture-in-transito/it/il-progetto"),
    ("22/5/2026", "Parole in transito, Episode 1 – ‘Appartenenza’. Contribution to the University of Bologna podcast developed from the workshop ‘AttraVerso la mia città’. The episode explores belonging through emotional maps, life stories and the everyday paths of migrant communities in Bologna.", "https://unibo.network.podigee.io/podcast/90078-parole-in-transito/1-appartenenza"),
    ("15/4/2026", "‘Mettere in scena la comprensione: Dario Fo e la comunicazione oltre le parole’. Contribution to the public programme celebrating the centenary of Dario Fo’s birth, DAMSLab / La Soffitta, University of Bologna.", "https://site.unibo.it/damslab/it/eventi/mistero-buffo-il-primo-miracolo-di-gesu-bambino-di-dario-fo-e-franca-rame"),
    ("5/6/2026", "‘Dal “bolognese doc” all’italiano “siento per cento”. Strategie lessicali di stereotipizzazione e posizionamento nella lingua parlata’, with Ana Pano Alamán. Public-facing contribution to the LILEC programme ‘Riscoprire, rielaborare, sovvertire il mondo’, on language, stereotypes and belonging.", "https://eventi.unibo.it/lilec-convegno-2026/programma"),
    ("27/9/2024", "‘Senti KIParla! Esplorare la diversità nell’italiano parlato’. Interactive activity for the European Researchers’ Night in Piazza Scaravilli, Bologna, with Eleonora Zucchini, Ludovica Pannitto and Claudia Borghetti. Games, quizzes and listening activities based on KIParla introduced the public to regional and social variation in spoken Italian.", "https://www.nottedeiricercatori-society.eu/eventi/piazza-scaravilli-senti-kiparla-esplorare-la-diversita-nellitaliano-parlato"),
    ("21/10/2022", "‘Corpus KIParla: un nuovo strumento per osservare l’italiano parlato e chi parla italiano’. Public-facing article for Linguisticamente explaining spoken-language corpora, sociolinguistic metadata and variation in interaction.", "https://www.linguisticamente.org/corpus-kiparla-un-nuovo-strumento-per-osservare-litaliano-parlato-e-chi-parla-italiano/"),
    ("4/3/2022", "‘Così vicino, così lontano! Prospettive sulla diversità linguistica’, with Silvia Ballarè. Two educational sessions for students at the Rolandino de’ Passaggeri lower-secondary school, Bologna."),
    ("2021", "‘Altro: quanto è diverso davvero l’Altro?’. Educational video module for the LILEC MOOC on words, languages and interculturality, designed for a broad non-specialist audience."),
]


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell(cell, text, url=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.add_run(text)
    if url:
        paragraph.add_run("\n")
        add_hyperlink(paragraph, url, url)


document = Document(SOURCE)
for section in document.sections:
    for paragraph in section.header.paragraphs:
        if "Last update" in paragraph.text:
            paragraph.text = paragraph.text.replace("30/6/2026", "13/8/2026")
publications_heading = next(p for p in document.paragraphs if p.text.strip() == "10. Publications")
invited_heading = next(p for p in document.paragraphs if p.text.strip() == "10) Invited talks and lectures")

# Reuse the existing two-column CV table design, including widths, borders and cell margins.
template = deepcopy(document.tables[10]._tbl)
while len(template.tr_lst) > len(ENTRIES):
    template.remove(template.tr_lst[-1])

heading = deepcopy(publications_heading._p)
for child in list(heading):
    if child.tag != qn("w:pPr"):
        heading.remove(child)
heading_paragraph = OxmlElement("w:r")
heading_properties = OxmlElement("w:rPr")
bold = OxmlElement("w:b")
size = OxmlElement("w:sz")
size.set(qn("w:val"), "22")
heading_properties.extend([bold, size])
heading_paragraph.append(heading_properties)
heading_text = OxmlElement("w:t")
heading_text.text = "10. Public engagement"
heading_paragraph.append(heading_text)
heading.append(heading_paragraph)

blank = OxmlElement("w:p")
publications_heading._p.addprevious(heading)
publications_heading._p.addprevious(blank)
publications_heading._p.addprevious(template)
publications_heading._p.addprevious(OxmlElement("w:p"))

new_table = next(t for t in document.tables if t._tbl is template)
for row, (date, description, *maybe_url) in zip(new_table.rows, ENTRIES):
    set_cell(row.cells[0], date)
    set_cell(row.cells[1], description, maybe_url[0] if maybe_url else None)

publications_heading.text = "11. Publications"
for run in publications_heading.runs:
    run.bold = True
    run.font.size = document.paragraphs[43].runs[0].font.size

invited_heading.text = "12. Invited talks and lectures"
for run in invited_heading.runs:
    run.bold = True
    run.font.size = document.paragraphs[43].runs[0].font.size

document.save(OUTPUT)
print(f"Saved {OUTPUT}")
