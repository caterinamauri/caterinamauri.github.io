from copy import deepcopy
from pathlib import Path
import re

from docx import Document


FILES = [Path("assets/Mauri_CV.docx"), Path("assets/Mauri_CV_tutto.docx")]

REPLACEMENTS = {
    "Ilaria Fiorentini": (
        "Post-doc in LEAdhoC — Linguistic expression of ad hoc categories."
    ),
    "Eugenio Goria": (
        "Post-doc in LEAdhoC — Linguistic expression of ad hoc categories."
    ),
    "Alessandra Barotto": (
        "Post-doc in LEAdhoC — Linguistic expression of ad hoc categories."
    ),
    "Eleonora Zucchini": (
        "Post-doc in DiverSIta — Diversity in Spoken Italian."
    ),
    "Giorgia Troiani": (
        "MSCA European Postdoctoral Fellowship, SPEAK — Convergence of syntactic, "
        "prosodic, and interactional units in conversation."
    ),
    "Elena Battaglia": (
        "SNSF Postdoc.Mobility, Towards an incremental semantics: A linguistic "
        "analysis of the emergence of categories in spoken interaction."
    ),
    "Antonia Russo": (
        "Research appointment at LILEC, Beyond Yes and No: Semantic and Pragmatic "
        "Dimensions of Responses in Spoken Italian."
    ),
}


def replace_project_description(paragraph, person, new_description):
    text = paragraph.text
    if person not in text:
        return False
    marker = " Current position:"
    if person == "Antonia Russo":
        if not any(label in text for label in ("forthcoming postdoctoral appointment", "Research appointment at LILEC")):
            return False
        prefix = text.split(" — ", 1)[0]
        new_text = (
            f"{prefix} — {new_description} Current position: Research Fellow, "
            "University of Bologna. https://www.unibo.it/sitoweb/antonia.russo3/en"
        )
    else:
        if marker not in text:
            return False
        prefix = text.split(" — ", 1)[0]
        suffix = marker + text.split(marker, 1)[1]
        suffix = re.sub(r"(https?://\S+)\s+\1", r"\1", suffix)
        new_text = f"{prefix} — {new_description}{suffix}"

    first_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)
    run = paragraph.add_run(new_text)
    if first_rpr is not None:
        run._r.insert(0, first_rpr)
    return True


def update(path):
    doc = Document(path)
    found = {name: 0 for name in REPLACEMENTS}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for name, description in REPLACEMENTS.items():
                        if replace_project_description(paragraph, name, description):
                            found[name] += 1

    missing = [name for name, count in found.items() if count != 1]
    if missing:
        raise RuntimeError(f"Unexpected match count in {path}: {found}")

    tmp = path.with_suffix(".tmp.docx")
    doc.save(tmp)
    Document(tmp)  # Re-open before replacing the canonical file.
    tmp.replace(path)
    return found


for file_path in FILES:
    print(file_path, update(file_path))
