#!/usr/bin/env python3
"""Add detailed postdoctoral supervision and newly organized events to the CV."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
CV = ROOT / "assets" / "Caterina-Mauri-CV.docx"


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    paragraph.text = text
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def find_table_row(document, fragment):
    for table in document.tables:
        for row in table.rows:
            if fragment in " ".join(cell.text for cell in row.cells):
                return table, row
    raise RuntimeError(f"Could not find table row containing {fragment!r}")


def insert_rows_after(table, anchor_row, entries):
    current = anchor_row._tr
    template = deepcopy(anchor_row._tr)
    for period, description in entries:
        new_row = deepcopy(template)
        current.addnext(new_row)
        current = new_row
        cells = new_row.findall(qn("w:tc"))
        if len(cells) < 2:
            raise RuntimeError("Expected a two-column CV table row")
        from docx.table import _Cell
        set_cell_text(_Cell(cells[0], table), period)
        set_cell_text(_Cell(cells[1], table), description)
        properties = new_row.find(qn("w:trPr"))
        if properties is not None:
            for height in list(properties.findall(qn("w:trHeight"))):
                properties.remove(height)


def main():
    document = Document(CV)

    # Replace the aggregate postdoc statement with a detailed, verifiable list.
    postdoc_table, postdoc_anchor = find_table_row(document, "Supervision of 7 PostDocs")
    postdoc_anchor_cells = postdoc_anchor.cells
    set_cell_text(postdoc_anchor_cells[0], "")
    set_cell_text(postdoc_anchor_cells[1], "Postdoctoral researchers supervised at the University of Bologna:")
    postdocs = [
        (
            "2016–2017",
            "Ilaria Fiorentini — LEAdhoC. Current position: Associate Professor of Linguistics, University of Pavia. https://unipv.unifind.cineca.it/resource/person/694724?language=it-IT",
        ),
        (
            "05/2016–02/2018",
            "Eugenio Goria — LEAdhoC. Current position: Associate Professor of Linguistics, University of Turin. https://unifind.unito.it/individual?uri=http%3A%2F%2Firises.unito.it%2Fresource%2Fperson%2F132686",
        ),
        (
            "10/2017–09/2021",
            "Alessandra Barotto — LEAdhoC. Current position: Tenure-track Assistant Professor of Linguistics, University of Insubria. https://uninsubria.unifind.cineca.it/resource/person/160735?language=en-US",
        ),
        (
            "03/2024–12/2025",
            "Eleonora Zucchini — DiverSIta. Current position: Assistant Professor, Department of Romance Languages and Literatures, Masaryk University. https://www.muni.cz/en/people/256684-eleonora-zucchini",
        ),
        (
            "01/2026–12/2027",
            "Giorgia Troiani — MSCA European Postdoctoral Fellowship, SPEAK. Current position: MSCA Postdoctoral Fellow, University of Bologna. https://www.unibo.it/sitoweb/giorgia.troiani/en",
        ),
        (
            "01/2026–12/2027",
            "Elena Battaglia — SNSF Postdoc.Mobility, Towards an incremental semantics: A linguistic analysis of the emergence of categories in spoken interaction. Current position: Postdoctoral Fellow, University of Bologna.",
        ),
        (
            "from 09/2026",
            "Antonia Russo — forthcoming postdoctoral appointment, University of Bologna. https://www.researchgate.net/profile/Antonia-Russo-2",
        ),
    ]
    insert_rows_after(postdoc_table, postdoc_anchor, postdocs)

    # Add events in reverse chronological order before the current first entry.
    events_table, first_event = find_table_row(document, "Relative Clauses Across and Within Languages")
    events = [
        (
            "27/6–2/7/2027",
            "Organizer, with Elena Battaglia, of the accepted panel ‘Towards an incremental semantics: the emergence of categories in spoken interaction’, 20th International Pragmatics Conference (IPC20), University of Helsinki, Finland. https://pragmatics.international/page/IPC20Helsinki",
        ),
        (
            "4–5/2/2027",
            "Organizer, with Giorgia Troiani, of the International Workshop ‘Interactional functions of syntax, syntactic functions of interaction’, MSCA-PF SPEAK, University of Bologna. https://site.unibo.it/msca_speak/en/agenda/workshop-functions-syntax-interaction",
        ),
        (
            "22/6/2026",
            "Organizer, with Giorgia Troiani, of the SPEAK seminar ‘Effects of language contact in oral corpora’, speaker Nina Dobrushina (Laboratoire Dynamique du Langage, CNRS Lyon), University of Bologna and online. Part of the seminar series ‘Identifying units of organization in spoken discourse: theory and methods’. https://site.unibo.it/msca_speak/en/agenda/talk-by-nina-dobrushina",
        ),
        (
            "9/4/2026",
            "Organizer, with Giorgia Troiani, of the SPEAK / CLUB seminar ‘Theoretical challenges in the annotation of spoken Kazakh and Italian: intonation, syntax, and interaction’, Experimental Laboratory, LILEC, University of Bologna. https://site.unibo.it/msca_speak/en/agenda/club-seminar-annotation-ius",
        ),
    ]
    # Insert before the first existing event while preserving the event-table format.
    first_xml = first_event._tr
    template = deepcopy(first_xml)
    previous = None
    for period, description in events:
        new_row = deepcopy(template)
        if previous is None:
            first_xml.addprevious(new_row)
        else:
            previous.addnext(new_row)
        previous = new_row
        cells = new_row.findall(qn("w:tc"))
        from docx.table import _Cell
        set_cell_text(_Cell(cells[0], events_table), period)
        set_cell_text(_Cell(cells[1], events_table), description)
        properties = new_row.find(qn("w:trPr"))
        if properties is not None:
            for height in list(properties.findall(qn("w:trHeight"))):
                properties.remove(height)

    document.save(CV)
    print(f"Updated {CV}")


if __name__ == "__main__":
    main()
