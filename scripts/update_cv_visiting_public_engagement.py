from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import _Row


DOCX = Path("assets/Caterina-Mauri-CV.docx")


def replace_cell_text(cell, text):
    """Replace visible text while retaining the formatting of the first run."""
    paragraph = cell.paragraphs[0]
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def find_row(document, needle):
    for table in document.tables:
        for row in table.rows:
            if needle in " | ".join(cell.text for cell in row.cells):
                return table, row
    raise RuntimeError(f"Row not found: {needle}")


def update_row(document, needle, date, description):
    _, row = find_row(document, needle)
    replace_cell_text(row.cells[0], date)
    replace_cell_text(row.cells[1], description)


def insert_public_engagement(document):
    table, anchor = find_row(document, "27/9/2024")
    # Clone an existing row so widths, borders, type and spacing remain unchanged.
    new_tr = deepcopy(anchor._tr)
    anchor._tr.addprevious(new_tr)
    new_row = _Row(new_tr, table)
    replace_cell_text(new_row.cells[0], "23/6/2025")
    replace_cell_text(
        new_row.cells[1],
        "‘La scena come osservatorio sociale’. Participant in the public workshop of "
        "the Homo Ludens programme, developed with director Eduardo Landim and "
        "Hospites Teatro, University of Bologna colleagues, students and third-sector "
        "organisations. The event explored theatre as a means of observing and "
        "reworking relational and cultural dynamics in educational, academic and "
        "social contexts. Centro Interculturale M. Zonarelli, Bologna.\n"
        "https://www.culturabologna.it/lang/ita/events/la-scena-come-osservatorio-sociale",
    )


LANDIM_DESCRIPTION = (
    "‘La scena come osservatorio sociale’. Participant in the public workshop of "
    "the Homo Ludens programme, developed with director Eduardo Landim and "
    "Hospites Teatro, University of Bologna colleagues, students and third-sector "
    "organisations. The event explored theatre as a means of observing and "
    "reworking relational and cultural dynamics in educational, academic and "
    "social contexts. Centro Interculturale M. Zonarelli, Bologna.\n"
    "https://www.culturabologna.it/lang/ita/events/la-scena-come-osservatorio-sociale"
)


def main():
    document = Document(DOCX)

    update_row(
        document,
        "17 April – 2 May /2026",
        "17 April–2 May 2026",
        "Guest Researcher, CNRS Laboratoire de Linguistique Formelle (LLF) and "
        "Université Paris Cité. Sponsor: Prof. Ira Noveck.",
    )
    update_row(
        document,
        "10-20 June /2022",
        "10–20 June 2022",
        "Visiting Scholar, Department of Linguistics, Stockholm University. "
        "Sponsor: Prof. Maria Koptjevskaja-Tamm.",
    )
    update_row(
        document,
        "August 2008",
        "February 2008",
        "Visiting Scholar, Max Planck Institute for Evolutionary Anthropology, "
        "Leipzig. Sponsor: Prof. Martin Haspelmath.",
    )
    update_row(
        document,
        "Aug.-Dec. /2005",
        "September–December 2005",
        "Visiting Student, Freie Universität Berlin. Sponsor: Prof. Ekkehard König.",
    )
    update_row(
        document,
        "Sept. /2004 – April /2005",
        "November 2004–February 2005",
        "Visiting Student, Universität Erfurt. Sponsor: Prof. Christian Lehmann.",
    )
    update_row(
        document,
        "Aug.-Sept. /2004",
        "August 2004",
        "Visiting Student, Max Planck Institute for Evolutionary Anthropology, "
        "Leipzig. Sponsor: Prof. Martin Haspelmath.",
    )

    if not any(
        "La scena come osservatorio sociale" in cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ):
        insert_public_engagement(document)
    else:
        update_row(
            document,
            "La scena come osservatorio sociale",
            "23/6/2025",
            LANDIM_DESCRIPTION,
        )

    document.save(DOCX)


if __name__ == "__main__":
    main()
