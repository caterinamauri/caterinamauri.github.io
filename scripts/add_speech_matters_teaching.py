from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import _Row


DOCX = Path("assets/Caterina-Mauri-CV.docx")


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def row_text(row):
    return " | ".join(cell.text for cell in row.cells)


def main():
    document = Document(DOCX)

    # The short course belongs under doctoral/advanced teaching, not invited talks.
    teaching_table = document.tables[13]
    if not any("Speech Matters" in row_text(row) for row in teaching_table.rows):
        anchor = teaching_table.rows[0]
        new_tr = deepcopy(anchor._tr)
        anchor._tr.addprevious(new_tr)
        new_row = _Row(new_tr, teaching_table)
        set_cell_text(new_row.cells[0], "19–22/5/2026")
        set_cell_text(
            new_row.cells[1],
            "Three-lecture course ‘Layers of cooperation: The interactional "
            "architecture of spoken language’, Speech Matters – 2nd Edition Spring "
            "School, Lake Como School of Advanced Studies, Villa del Grumello, Como. "
            "Lectures held on 19, 21 and 22 May 2026. "
            "https://spma.lakecomoschool.org/program/",
        )

    for table_index, table in enumerate(document.tables):
        if table_index == 13:
            continue
        for row in list(table.rows):
            if (
                "Layers of cooperation: The interactional architecture of spoken language"
                in row_text(row)
            ):
                table._tbl.remove(row._tr)

    document.save(DOCX)


if __name__ == "__main__":
    main()
