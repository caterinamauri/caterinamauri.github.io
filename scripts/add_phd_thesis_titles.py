#!/usr/bin/env python3
"""Add detailed PhD mentoring and final-examination thesis titles to both CVs."""

from copy import deepcopy
from pathlib import Path
import tempfile

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Row


ROOT = Path(__file__).resolve().parents[1]
FILES = [ROOT / "assets" / "Mauri_CV.docx", ROOT / "assets" / "Mauri_CV_tutto.docx"]

PHD_GROUPS = [
    ("PhD thesis supervisor: 1", [
        "Alessandra Barotto — Exemplification and Categorization: The Case of Japanese.",
    ]),
    ("PhD thesis co-supervisor: 3", [
        "Maria Cristina Lo Baido — The Comment Clause in Present-day Italian: Forms, Functions, Directionalities.",
        "Antonia Russo — L’espressione linguistica della sostituzione: uno studio sull’italiano tra sincronia e diacronia.",
        "Kristýna Lorenzová — The Development of Discourse Markers in the Interlanguages of Slavic Learners of Italian as a Non-Mother Tongue: A Longitudinal Study.",
    ]),
    ("International doctoral mentor: 1", [
        "Elena Battaglia — Evidentiality and Interaction in Italian.",
    ]),
    ("Appointed doctoral thesis reader (ongoing advisory role): 7", [
        "Antonio Bianco — L’uso dello humor nella campagna elettorale 2022.",
        "Sara Gemelli — Gender-Based Violence in Media Representations in Italy: Critical Discourse Analysis, Framing, and Perception Studies.",
        "Tanja Trebucchi — I marcatori pragmatici nel discorso politico italiano: (inter)soggettività, percezione e persuasione.",
        "Jessica Katiuscia Ivani — The Morphosyntax of Number Systems: A Cross-Linguistic Study.",
        "Eugenio Goria — Il discorso bilingue a Gibilterra: l’emergere di schemi regolari nella commutazione di codice.",
        "Luigi Talamo — Nominalizations of Property Concepts: Evidence from Italian.",
        "Simone Mattiola — Typology of Pluractional Constructions in the Languages of the World.",
    ]),
]

COMMISSIONS = [
    ("2021", "Marco Favaro — Pragmatic Markers in Italian: Four Case Studies on Illocutive Functions of Adverbs and Sociolinguistic Variation."),
    ("2024", "Roberta Cicchirillo — Aspetti prosodici, sintattici e semantici per una definizione multidimensionale delle costruzioni a lista."),
    ("2024", "Roxanne Holly Padley — A Corpus-Based Analysis of Healthcare Communication and Interaction through Cosmetic Surgery Discourses."),
    ("2026", "Tommaso Lamarra — Language-Mediated Abstraction: Conceptual Concreteness and Categorical Specificity in Language Processing."),
]


def row_text(row):
    return " | ".join(cell.text for cell in row.cells)


def set_paragraph_text_preserving_format(paragraph, text):
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def set_cell(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text_preserving_format(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def fill_row(row, left, right):
    cells = row.cells
    set_cell(cells[0], left)
    if cells[0]._tc is not cells[-1]._tc:
        set_cell(cells[-1], right)


def cloned_row(template, left, right):
    tr = deepcopy(template._tr)
    # Temporarily wrap the cloned row so python-docx exposes its cells/runs.
    template._tr.addnext(tr)
    new_row = _Row(tr, template._parent)
    fill_row(new_row, left, right)
    return tr


def update(path):
    doc = Document(path)
    table = next(
        t for t in doc.tables
        if any("PhD students supervised" in row_text(r) or "Doctoral supervision and continuing" in row_text(r) for r in t.rows)
    )
    rows = table.rows
    postdoc_rows = [r for r in rows if "Postdoctoral researchers supervised" in row_text(r)]
    postdoc_heading = postdoc_rows[0]
    for duplicate in postdoc_rows[1:]:
        duplicate._tr.getparent().remove(duplicate._tr)
    rows = table.rows
    phd_heading = next(
        r for r in rows
        if "PhD students supervised" in row_text(r) or "Doctoral supervision and continuing" in row_text(r)
    )
    detail_template = next(
        (r for r in rows if "PhD Program in Linguistic Sciences" in row_text(r)),
        next(r for r in rows if "Post-doc in" in row_text(r)),
    )

    set_cell(phd_heading.cells[0], "-- Doctoral supervision and continuing thesis advisory roles: 12 total")

    # Remove the old aggregate Pavia/Bergamo and USI rows.
    current = phd_heading._tr.getnext()
    while current is not None and current is not postdoc_heading._tr:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    insert_before = postdoc_heading._tr
    for group_heading, entries in PHD_GROUPS:
        insert_before.addprevious(cloned_row(phd_heading, group_heading, ""))
        for entry in entries:
            insert_before.addprevious(cloned_row(detail_template, "", entry))

    commission_heading = cloned_row(phd_heading, "-- Final PhD examination committees: 4 total", "")
    insert_before.addprevious(commission_heading)
    for left, right in COMMISSIONS:
        insert_before.addprevious(cloned_row(detail_template, left, right))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=path.parent) as tmp:
        tmp_path = Path(tmp.name)
    doc.save(tmp_path)
    tmp_path.replace(path)


if __name__ == "__main__":
    for file in FILES:
        update(file)
        print(f"Updated {file}")
