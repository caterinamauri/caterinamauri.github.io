#!/usr/bin/env python3
"""Move thesis lists into section 8 and add native Word numbering."""

from pathlib import Path
import re
import tempfile
import unicodedata
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
FILES = [ROOT / "assets" / "Mauri_CV.docx", ROOT / "assets" / "Mauri_CV_tutto.docx"]
XLSX_PRIMARY = Path("/Users/caterinamauri/Downloads/elenco1.xlsx")
XLSX_SECONDARY = Path("/Users/caterinamauri/Downloads/elenco2.xlsx")


def next_id(elements, attr):
    values = [int(el.get(qn(attr))) for el in elements if el.get(qn(attr), "").isdigit()]
    return max(values, default=0) + 1


def add_decimal_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_id = next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = next_id(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    lvl.append(text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "right")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id=None):
    ppr = paragraph._p.get_or_add_pPr()
    for old in list(ppr.findall(qn("w:numPr"))):
        ppr.remove(old)
    if num_id is None:
        return
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numpr.append(ilvl)
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(numid)
    ppr.append(numpr)


def replace_text_preserve_first_run(paragraph, text):
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def normalize(value):
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode().casefold()
    return re.sub(r"[^a-z0-9]+", " ", value).strip()


def xlsx_theses(path):
    ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    records = []
    with zipfile.ZipFile(path) as archive:
        shared = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            shared = ["".join(t.text or "" for t in si.findall(".//m:t", ns)) for si in root.findall("m:si", ns)]
        sheet = ET.fromstring(archive.read("xl/worksheets/sheet1.xml"))
        for row in sheet.findall(".//m:sheetData/m:row", ns):
            values = {}
            for cell in row.findall("m:c", ns):
                col = re.match(r"[A-Z]+", cell.get("r")).group()
                value_node = cell.find("m:v", ns)
                value = "" if value_node is None else (value_node.text or "")
                if cell.get("t") == "s" and value:
                    value = shared[int(value)]
                elif cell.get("t") == "inlineStr":
                    value = "".join(t.text or "" for t in cell.findall(".//m:t", ns))
                values[col] = re.sub(r"\s+", " ", value).strip()
            year_match = re.search(r"(?:19|20)\d{2}", values.get("M", ""))
            if values.get("G") and year_match:
                surname = values.get("B", "").strip()
                given = values.get("C", "").strip()
                name = ", ".join(part for part in (surname.title(), given.title()) if part)
                records.append({
                    "title": values["G"],
                    "year": year_match.group(),
                    "degree": values.get("I", ""),
                    "student": name,
                })
    return records


def pavia_theses():
    course = "LINGUISTICA TEORICA, APPLICATA E DELLE LINGUE MODERNE [05409]"
    # The public archive exposes the academic year, not the exact defence date;
    # the calendar year below is the concluding year of that academic year.
    return [
        {"student": "Laboranti, Sara Camilla", "title": "Categorizzare con liste ed esempi: uno studio condotto attraverso l'eye tracking", "year": "2015", "degree": course},
        {"student": "Cavanna, Gloria", "title": "I connettivi ipotetico-presupposizionali dell'italiano contemporaneo: definizione e proposta di analisi", "year": "2013", "degree": course},
        {"student": "Runello, Chiara", "title": "Il netspeak dei fandom - un'analisi linguistica", "year": "2015", "degree": course},
        {"student": "Pigozzi, Cecilia", "title": "Le funzioni dei complimenti", "year": "2013", "degree": course},
        {"student": "Lo Baido, Maria Cristina", "title": "Le funzioni dell'esemplificazione in italiano: tra cognizione e discorso", "year": "2015", "degree": course},
        {"student": "Manzinello, Valeria", "title": "Strutture comparative a base nominale: graduabilità e categorie cognitive", "year": "2016", "degree": course},
        {"student": "Alè, Francesco Antonio", "title": "Un corpus per segnare: corpus linguistics nella didattica e nell'apprendimento della lingua dei segni", "year": "2012", "degree": course},
        {"student": "Tamburini, Paolo", "title": "Dentro la barzelletta: moderni approcci linguistico-cognitivi allo studio di storielle divertenti", "year": "2016", "degree": course},
        {"student": "Baesso, Martina", "title": "Il code switching: variazione di codice in madrelingua cinesi residenti in Italia", "year": "2015", "degree": course},
    ]


def thesis_metadata():
    records = xlsx_theses(XLSX_PRIMARY) + xlsx_theses(XLSX_SECONDARY) + pavia_theses()
    return {normalize(record["title"]): record for record in records}


def paragraphs_between(document, start_text, end_text):
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start_text)
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == end_text)
    return paragraphs[start + 1:end]


def number_publications(document):
    pub_num = add_decimal_numbering(document)
    category_names = {
        "Monographs",
        "Edited volumes and special issues",
        "Articles in peer-reviewed journals",
        "Papers and chapters within edited volumes",
    }
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "10. Publications")
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Linguistic resources")
    count = 0
    for p in paragraphs[start + 1:end]:
        text = p.text.strip()
        if not text or text in category_names:
            set_numbering(p, None)
            continue
        cleaned = re.sub(r"^\d+\.\s*", "", p.text, count=1)
        if cleaned != p.text:
            replace_text_preserve_first_run(p, cleaned)
        set_numbering(p, pub_num)
        count += 1
    return count


def move_and_number_theses(document):
    paragraphs = document.paragraphs
    thesis_heading = next((p for p in paragraphs if p.text.strip() == "11. Supervised BA and MA theses"), None)
    if thesis_heading is None:
        return 0
    invited = next(p for p in paragraphs if p.text.strip() == "12. Invited talks and lectures")
    section9 = next(p for p in paragraphs if p.text.strip() == "9. Public engagement")
    body = document.element.body

    # Gather the paragraphs between the temporary thesis heading and invited talks.
    children = list(body)
    start = children.index(thesis_heading._p)
    end = children.index(invited._p)
    thesis_elements = [el for el in children[start + 1:end] if el.tag == qn("w:p")]
    thesis_heading._p.getparent().remove(thesis_heading._p)

    labels = {
        "University of Bologna — Primary supervisor",
        "University of Pavia — Primary supervisor",
        "University of Bologna — Co-supervisor",
    }
    thesis_num = add_decimal_numbering(document)
    metadata = thesis_metadata()
    count = 0
    for element in thesis_elements:
        text = "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()
        if not text:
            element.getparent().remove(element)
            continue
        # Moving an existing element before section 9 also removes it from its old position.
        body.insert(list(body).index(section9._p), element)
        paragraph = next(p for p in document.paragraphs if p._p is element)
        if text in labels:
            set_numbering(paragraph, None)
        else:
            title = re.sub(r"\s+\((?:19|20)\d{2}(?:/(?:19|20)\d{2})?\)\..*$", "", text).strip()
            record = metadata.get(normalize(title))
            if record is None:
                raise RuntimeError(f"No student metadata found for thesis: {title}")
            replace_text_preserve_first_run(
                paragraph,
                f"{record['student']}. {record['title']} ({record['year']}). {record['degree']}",
            )
            set_numbering(paragraph, thesis_num)
            count += 1

    replace_text_preserve_first_run(invited, "11. Invited talks and lectures")
    return count


def update_supervision_summary(document):
    target = next(
        table for table in document.tables
        if "MA and BA students supervised at the Universities of Pavia and Bologna" in " ".join(cell.text for row in table.rows for cell in row.cells)
    )
    replacements = {
        "-- MA and BA students supervised at the Universities of Pavia and Bologna:":
            "-- BA and MA theses supervised or co-supervised: 99 total (54 BA, 45 MA):",
        "Dept. of Human Studies, Univ. Pavia. Supervision of 15 MA students":
            "Dept. of Human Studies, Univ. Pavia. Supervision of 9 MA theses.",
        "Dept. of Modern languages, literatures and cultures, Univ. Bologna. Supervision of 89 graduate students (30 MA, 59 BA).":
            "Dept. of Modern Languages, Literatures and Cultures, Univ. Bologna. Supervision or co-supervision of 90 theses (36 MA, 54 BA).",
    }
    changed = set()
    for row in target.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                for old, new in replacements.items():
                    if text == old:
                        replace_text_preserve_first_run(paragraph, new)
                        changed.add(old)
    if set(replacements) != changed:
        missing = set(replacements) - changed
        raise RuntimeError(f"Could not update supervision summary: {missing}")


def save_atomic(document, path):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=path.parent) as handle:
        temp = Path(handle.name)
    document.save(temp)
    # Confirm python-docx can reopen the result before replacement.
    Document(temp)
    temp.replace(path)


def main():
    for path in FILES:
        document = Document(path)
        publication_count = number_publications(document)
        thesis_count = move_and_number_theses(document)
        update_supervision_summary(document)
        save_atomic(document, path)
        print(path.name, publication_count, "publications;", thesis_count, "theses")


if __name__ == "__main__":
    main()
